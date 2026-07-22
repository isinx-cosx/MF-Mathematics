# -*- coding: utf-8 -*-
"""MF-Mathematics 用户手册 Word 文档生成器。"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color):
    """设置单元格背景色。"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    tcPr.append(shading)


def add_image_placeholder(doc, caption, width_inches=5.5, height_inches=2.5):
    """添加图片占位框 + 题注。"""
    # 占位框（用单单元格表格模拟虚线框）
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, 'F5F5F4')

    # 设置边框虚线
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'dashed')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:space'), '4')
        border.set(qn('w:color'), 'A8A29E')
        borders.append(border)
    tcPr.append(borders)

    # 占位文字
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('🖼')
    run.font.size = Pt(28)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(caption)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x78, 0x71, 0x6C)
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(f'[ 截图后替换此占位框 · 建议宽度 {int(width_inches * 2.54):.0f} mm ]')
    r3.font.size = Pt(8)
    r3.font.color.rgb = RGBColor(0xA8, 0xA2, 0x9E)

    # 题注
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    cap_run.font.size = Pt(9)
    cap_run.font.color.rgb = RGBColor(0x78, 0x71, 0x6C)
    cap_run.italic = True

    doc.add_paragraph()  # 间距


def tip_box(doc, text):
    """蓝色提示框。"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run('💡 提示：')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)


def warn_box(doc, text):
    """黄色警告框。"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run('⚠️ 注意：')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xF5, 0x9E, 0x0B)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)


def add_table(doc, headers, rows):
    """添加格式化表格。"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    # 数据
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.cell(ri + 1, ci)
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════
def build_manual(output_path: str):
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.5
    # 设置中文字体回退
    rPr = style.element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    rPr.insert(0, rFonts)

    # ═══════════════════════════════════════════════════════════════
    #  封面
    # ═══════════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('MF-Mathematics')
    r.font.size = Pt(36)
    r.bold = True
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('多功能数学计算与可视化平台')
    r2.font.size = Pt(16)
    r2.font.color.rgb = RGBColor(0x78, 0x71, 0x6C)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run('用户手册  ·  V1.0.0  ·  2026 年 7 月')
    r3.font.size = Pt(12)
    r3.font.color.rgb = RGBColor(0xA8, 0xA2, 0x9E)

    add_image_placeholder(doc, '应用图标', 2.5, 2.5)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run('MF-Vis-Science 工作室')
    r4.font.size = Pt(11)
    r4.font.color.rgb = RGBColor(0x78, 0x71, 0x6C)
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    #  目录页
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('目  录', level=1)
    toc_items = [
        '第 1 章 · 软件概述',
        '第 2 章 · 安装与启动',
        '第 3 章 · 主界面概览',
        '第 4 章 · 计算功能（14 个模块）',
        '第 5 章 · 绘图功能（7 个模式）',
        '第 6 章 · 用户账户系统',
        '第 7 章 · AI 智能助手',
        '第 8 章 · 键盘面板',
        '第 9 章 · 设置与主题',
        '第 10 章 · 常见问题',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(1)
        p.runs[0].font.size = Pt(11)
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    #  第 1 章
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('第 1 章 · 软件概述', level=1)

    doc.add_heading('1.1 产品简介', level=2)
    doc.add_paragraph(
        'MF-Mathematics 是一款面向数学学习、科研与工程计算的桌面端多功能数学平台。'
        '软件集成符号计算、数值分析、函数绘图、分形探索、AI 辅助等核心能力，'
        '覆盖从初等数学到高等数学的广泛领域。'
    )

    doc.add_heading('1.2 核心功能', level=2)
    add_table(doc,
        ['类别', '内容', '数量'],
        [['计算模块', '基础运算、代数、微积分、解析几何、数列、线代、概率统计、数值分析、数论、实分析、泛函分析、复分析、代数拓扑、测度论', '14'],
         ['绘图模块', '普通 2D、极坐标、3D 曲面、复数映射、向量场、任意做图、分形探索', '7'],
         ['AI 助手', '自然语言输入数学问题，逐步推理与解答', '—'],
         ['用户系统', '在线注册、登录、邮箱验证、余额管理', '—'],
         ['主题', '亮色（Slate）+ 暗色（Catppuccin Mocha）双主题', '2']]
    )

    doc.add_heading('1.3 系统要求', level=2)
    add_table(doc,
        ['项目', '要求'],
        [['操作系统', 'Windows 10 / 11（64 位）'],
         ['内存', '建议 4 GB 及以上'],
         ['硬盘空间', '约 300 MB'],
         ['网络', '登录/注册/AI 功能需要互联网连接']]
    )

    # ═══════════════════════════════════════════════════════════════
    #  第 2 章
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('第 2 章 · 安装与启动', level=1)

    doc.add_heading('2.1 获取安装包', level=2)
    doc.add_paragraph('访问 MF-Mathematics 官方网站下载最新安装包：')
    doc.add_paragraph('https://www.mvs-studio.com/download')
    add_image_placeholder(doc, '图 2-1：官网下载页面', 5.5, 2.2)

    doc.add_heading('2.2 安装步骤', level=2)
    for i, step in enumerate([
        '双击 Multifunctional-Mathematics-Setup.exe',
        '选择安装语言（简体中文 / English）',
        '阅读并接受许可协议',
        '选择安装目录（默认 C:\\Program Files\\MF-Mathematics）',
        '勾选"创建桌面快捷方式"',
        '点击"安装"，等待完成',
    ], 1):
        doc.add_paragraph(f'{i}. {step}')
    add_image_placeholder(doc, '图 2-2：安装向导', 4.0, 2.5)

    doc.add_heading('2.3 启动与更新', level=2)
    doc.add_paragraph(
        '双击桌面快捷方式或从开始菜单启动。软件启动时会自动检测新版本，如有更新将弹出提示对话框。'
    )
    add_image_placeholder(doc, '图 2-3：版本更新提示', 4.5, 1.8)

    # ═══════════════════════════════════════════════════════════════
    #  第 3 章
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('第 3 章 · 主界面概览', level=1)
    add_image_placeholder(doc, '图 3-1：主界面全貌（亮色主题）· 请标注 ①~⑥ 区域编号', 5.5, 3.5)

    doc.add_heading('3.1 界面布局', level=2)
    add_table(doc,
        ['编号', '区域', '说明'],
        [['①', '菜单栏', '文件、编辑、视图、主题切换、帮助'],
         ['②', '工具栏', '计算/绘图模式切换、子导航、登录按钮'],
         ['③', '子导航栏', '当前模式下的二级功能标签页'],
         ['④', '工作区', '计算块 / 绘图画布 / 数学显示区'],
         ['⑤', '键盘面板', '虚拟数学键盘（可折叠）'],
         ['⑥', '状态栏', '就绪信息、当前用户、MF-Vis-Science 品牌']]
    )

    doc.add_heading('3.2 菜单栏', level=2)
    add_table(doc,
        ['菜单', '功能项'],
        [['文件', '新建计算、保存工作区、导出图片、退出'],
         ['编辑', '撤销、重做、复制、粘贴、清空工作区'],
         ['视图', '切换亮色/暗色主题、显示/隐藏键盘面板'],
         ['工具', 'AI 助手、搜索面板、计算历史'],
         ['帮助', '用户手册、关于、版本检测']]
    )

    doc.add_heading('3.3 工具栏与模式切换', level=2)
    doc.add_paragraph(
        '工具栏左侧显示当前模式名称，点击可在计算模式和绘图模式之间切换。'
        '右侧包含 AI 助手入口、用户登录按钮和帮助按钮。'
    )
    add_image_placeholder(doc, '图 3-2：工具栏', 5.5, 0.6)

    # ═══════════════════════════════════════════════════════════════
    #  第 4 章
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('第 4 章 · 计算功能', level=1)
    doc.add_paragraph(
        '软件提供 14 个计算模块，覆盖从基础算术到高等数学的完整体系。'
        '每个模块采用统一的 CalcBlock（计算块）交互模式：输入数学表达式 → 点击计算 → 查看分步结果。'
    )

    calc_modules = [
        ('4.1 基础运算', '四则运算、乘方、开方、阶乘、对数等基本数学运算。支持表达式连续输入和即时求值。', '图 4-1：基础运算界面'),
        ('4.2 代数计算', '多项式展开、因式分解、方程求解、不等式求解、表达式化简。支持一元和多元代数运算。', '图 4-2：代数计算界面'),
        ('4.3 微积分', '求导（高阶导数）、不定积分、定积分、极限计算、级数展开（泰勒/傅里叶）。符号推导与数值计算双模式。', '图 4-3：微积分界面'),
        ('4.4 解析几何', '直线、圆、椭圆、双曲线、抛物线的方程分析。交点计算、切线方程、焦点/准线等几何性质求解。', None),
        ('4.5 数列', '等差数列、等比数列、递推数列的通项公式与求和。数列极限分析、收敛性判断。', None),
        ('4.6 线性代数', '矩阵运算（加减乘、转置、求逆）、行列式、特征值与特征向量、线性方程组求解、二次型标准化。', None),
        ('4.7 概率论与数理统计', '概率分布（二项/泊松/正态/指数等）、期望与方差、假设检验、ANOVA、回归分析、时间序列。', None),
        ('4.8 数值分析', '数值积分、常微分方程求解（欧拉/龙格-库塔）、插值与拟合、非线性方程求根（牛顿法/二分法）、误差分析。', None),
        ('4.9 数论', '素数筛法、最大公约数/最小公倍数、模运算、同余方程、连分数、算术函数（欧拉 φ、莫比乌斯 μ）。', None),
        ('4.10 实分析', '实数完备性、序列极限（ε-N）、函数极限（ε-δ）、可微性判定、黎曼积分。', None),
        ('4.11 泛函分析', '赋范空间、内积空间、线性算子、对偶空间、谱理论、核心定理（Hahn-Banach、开映射、闭图像）。', None),
        ('4.12 复分析', '全纯函数判定、柯西-黎曼方程、复积分、留数定理、共形映射、黎曼 ζ 函数。', None),
        ('4.13 代数拓扑', '同伦群、同调群、上同调环、度理论与不动点、持续同调。', None),
        ('4.14 测度论', 'σ-代数构造、测度构造、可测函数、勒贝格积分、收敛定理、乘积测度、概率测度。', None),
    ]

    for title, desc, img_caption in calc_modules:
        doc.add_heading(title, level=2)
        doc.add_paragraph(desc)
        if img_caption:
            add_image_placeholder(doc, img_caption, 5.5, 2.0)

    tip_box(doc, '每个计算模块均支持多块管理——点击"添加计算块"可并行进行多个独立计算，块之间互不影响。')
    add_image_placeholder(doc, '图 4-4：多计算块并行工作', 5.5, 2.2)

    # ═══════════════════════════════════════════════════════════════
    #  第 5 章
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('第 5 章 · 绘图功能', level=1)
    doc.add_paragraph(
        '软件提供 7 个绘图模式，基于 Matplotlib 渲染，支持交互式缩放、平移、导出。'
    )

    plot_modules = [
        ('5.1 普通 2D 绘图', '绘制一元函数曲线 y = f(x)。支持多函数叠加、颜色自定义、坐标轴范围调整、网格线开关。', '图 5-1：2D 多函数绘图'),
        ('5.2 极坐标绘图', '绘制极坐标曲线 r = f(θ)。支持玫瑰线、心形线、螺线等经典极坐标图形。', '图 5-2：极坐标绘图'),
        ('5.3 3D 曲面绘图', '绘制二元函数曲面 z = f(x, y)。支持旋转视角（鼠标拖拽）、缩放、配色方案切换。', '图 5-3：3D 曲面绘图'),
        ('5.4 复数映射可视化', '将复变函数 w = f(z) 作用于网格或区域，直观展示复平面的映射效果。', None),
        ('5.5 向量场', '绘制二维向量场 F(x, y) = (P, Q)。支持流线、箭头密度调节。', None),
        ('5.6 任意做图', '自由几何构造：点、线段、圆、多边形、角度标记。支持撤销/重做、几何约束。', '图 5-4：任意几何做图'),
        ('5.7 分形探索', '生成经典分形图形：Mandelbrot 集、Julia 集、Newton 分形。支持无限缩放（鼠标滚轮）和色彩映射调节。', '图 5-5：Mandelbrot 分形'),
    ]

    for title, desc, img_caption in plot_modules:
        doc.add_heading(title, level=2)
        doc.add_paragraph(desc)
        if img_caption:
            add_image_placeholder(doc, img_caption, 5.5, 3.0)

    # ═══════════════════════════════════════════════════════════════
    #  第 6 章
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('第 6 章 · 用户账户系统', level=1)

    doc.add_heading('6.1 登录', level=2)
    doc.add_paragraph('点击工具栏"登录"按钮，打开登录/注册对话框。输入用户名和密码，点击"登录"。')
    add_image_placeholder(doc, '图 6-1：登录页面', 3.5, 3.0)

    doc.add_heading('6.2 注册', level=2)
    for i, step in enumerate([
        '切换到"注册"标签页',
        '输入用户名（3-20 字符，字母/数字/下划线）',
        '输入邮箱 → 点击"发送验证码"',
        '查收邮件，输入 6 位验证码',
        '设置密码（至少 8 位，含字母和数字）',
        '确认密码 → 点击"注册"',
    ], 1):
        doc.add_paragraph(f'{i}. {step}')
    add_image_placeholder(doc, '图 6-2：注册页面', 3.5, 3.6)

    warn_box(doc, '验证码有效期为 5 分钟，60 秒内不可重复发送。注册成功后自动登录并赠送 10 次免费额度。')

    doc.add_heading('6.3 邮箱验证（独立流程）', level=2)
    doc.add_paragraph('如果注册时未提供验证码，系统将发送验证邮件。用户需在验证页面输入验证码完成激活。')
    add_image_placeholder(doc, '图 6-3：邮箱验证页面', 3.5, 2.8)

    doc.add_heading('6.4 登录后状态', level=2)
    doc.add_paragraph('登录成功后，工具栏按钮显示"用户名 (余额: N)"，状态栏同步更新。点击该按钮可登出。')
    add_image_placeholder(doc, '图 6-4：已登录状态栏', 5.5, 0.4)

    # ═══════════════════════════════════════════════════════════════
    #  第 7 章
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('第 7 章 · AI 智能助手', level=1)
    doc.add_paragraph(
        'AI 助手支持自然语言输入数学问题，提供逐步推理和解答。由 DeepSeek / OpenAI 大语言模型驱动。'
    )

    doc.add_heading('7.1 打开 AI 助手', level=2)
    doc.add_paragraph('工具栏点击"AI"按钮，或菜单栏 → 工具 → AI 助手。')

    doc.add_heading('7.2 使用方式', level=2)
    for i, step in enumerate([
        '在输入框输入数学问题（如"求 x² + 3x + 2 = 0 的根"）',
        '点击发送或按 Enter 键',
        'AI 返回分步推理和最终答案',
    ], 1):
        doc.add_paragraph(f'{i}. {step}')

    doc.add_heading('7.3 配置 API Key', level=2)
    for i, step in enumerate([
        '菜单栏 → 工具 → AI 设置',
        '选择服务商（DeepSeek / OpenAI）',
        '填入 API Key',
        '点击保存',
    ], 1):
        doc.add_paragraph(f'{i}. {step}')
    add_image_placeholder(doc, '图 7-1：AI 助手对话框', 4.5, 3.0)
    tip_box(doc, 'AI 助手会自动读取当前工作区的计算内容作为上下文，使回答更加精准。')

    # ═══════════════════════════════════════════════════════════════
    #  第 8 章
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('第 8 章 · 键盘面板', level=1)
    doc.add_paragraph('内置虚拟数学键盘，提供常用数学符号和函数的快速输入。')

    doc.add_heading('8.1 面板布局', level=2)
    add_table(doc,
        ['区域', '内容'],
        [['数字区', '0-9、小数点、正负号'],
         ['运算符', '+、−、×、÷、^、√、='],
         ['函数区', 'sin、cos、tan、log、ln、exp、abs'],
         ['符号区', 'π、e、i、∞、(、)、,'],
         ['微积分', 'd/dx、∫、lim、Σ'],
         ['希腊字母', 'α、β、γ、θ、λ、μ、σ、φ、ω']]
    )
    add_image_placeholder(doc, '图 8-1：键盘面板', 5.5, 1.5)

    doc.add_heading('8.2 显示/隐藏', level=2)
    doc.add_paragraph('菜单栏 → 视图 → 显示/隐藏键盘面板，或点击工作区底部的折叠按钮。')

    # ═══════════════════════════════════════════════════════════════
    #  第 9 章
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('第 9 章 · 设置与主题', level=1)

    doc.add_heading('9.1 主题切换', level=2)
    doc.add_paragraph('软件内置两套完整主题：')
    add_table(doc,
        ['主题', '风格', '适用场景'],
        [['亮色（Slate）', '白底深色文字，蓝调强调', '日间使用'],
         ['暗色（Catppuccin Mocha）', '深紫黑底浅色文字', '夜间/护眼']]
    )
    doc.add_paragraph('菜单栏 → 视图 → 切换主题，或使用快捷键 Ctrl+T。')
    add_image_placeholder(doc, '图 9-1：暗色主题', 5.5, 3.5)

    doc.add_heading('9.2 设置对话框', level=2)
    doc.add_paragraph('菜单栏 → 工具 → 设置，可配置：')
    for item in ['默认绘图范围', '数值精度', 'AI 服务商与模型', '语言偏好']:
        doc.add_paragraph(f'• {item}')
    add_image_placeholder(doc, '图 9-2：设置对话框', 4.0, 2.5)

    # ═══════════════════════════════════════════════════════════════
    #  第 10 章
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('第 10 章 · 常见问题', level=1)

    faqs = [
        ('Q1：启动时提示"无法连接到服务器"？',
         '登录和 AI 功能需要网络连接。本地计算和绘图功能不受影响，可离线使用。'),
        ('Q2：忘记密码怎么办？',
         '目前请通过注册邮箱联系管理员重置。后续版本将增加自助找回密码功能。'),
        ('Q3：计算结果显示"计算超时"？',
         '复杂表达式可能超出计算时间限制。尝试简化表达式或分步计算。'),
        ('Q4：3D 图形无法旋转？',
         '在 3D 画布区域按住鼠标左键拖拽即可旋转视角，滚轮缩放。'),
        ('Q5：如何导出计算结果？',
         '菜单栏 → 文件 → 导出。支持将当前工作区导出为图片或 LaTeX 格式。'),
        ('Q6：验证码未收到？',
         '检查垃圾邮件箱。确保输入的邮箱地址正确。60 秒后可重新发送。'),
        ('Q7：如何切换语言？',
         '当前版本为简体中文。多语言支持将在后续版本中添加。'),
    ]

    for q, a in faqs:
        doc.add_heading(q, level=3)
        doc.add_paragraph(a)

    # ── 尾页 ──
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('MF-Mathematics V1.0.0 用户手册')
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xA8, 0xA2, 0x9E)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('MF-Vis-Science 工作室  © 2026')
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0xA8, 0xA2, 0x9E)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run('https://www.mvs-studio.com')
    r3.font.size = Pt(10)
    r3.font.color.rgb = RGBColor(0xA8, 0xA2, 0x9E)

    # 保存
    doc.save(output_path)


if __name__ == '__main__':
    output = os.path.join(
        os.path.dirname(os.path.abspath(__file__)), '..',
        'dist', 'copyright', 'MF-Mathematics_用户手册_V1.0.0.docx'
    )
    os.makedirs(os.path.dirname(output), exist_ok=True)
    build_manual(output)
    print(f'用户手册已生成: {output}')
    print(f'大小: {os.path.getsize(output) / 1024:.0f} KB')
